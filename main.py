"""
Contract Compliance Checker
Compares a signed contract against a master template to flag
missing clauses, altered definitions, and unfilled placeholders.
"""

from docx import Document


def extract_text(file_path):
    """Extract raw paragraph text from a .docx file."""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paragraphs

def is_heading(paragraph):
    # .style.name gives you the user-facing name of the style
    style_name = paragraph.style.name
    
    if style_name.startswith("Heading"):
        print(f"Found a heading: '{paragraph.text}' (Style: {style_name})")
        return True
    
    # 1. Check length
    if len(paragraph.text) < 60:
        # 2. Check casing (Either ALL CAPS or Title Case/Capitalized)
        is_correct_case = paragraph.text.isupper() or paragraph.text.istitle()

        # 3. Check if any part of the paragraph is bold
        is_bold = any(run.bold for run in paragraph.runs)

        # Combine the logic
        if is_correct_case and is_bold:
            print(f"Found a potential heading: '{paragraph.text}'")
            return True

    return False

def split_into_sections(file_path):
    """
    Group a flat list of paragraphs into sections, keyed by heading.
    """
    doc = Document(file_path)
    sections = {}
    current_heading = None

    for p in doc.paragraphs:
        if is_heading(p):
            current_heading = p.text
            sections[current_heading] = []
        else:
            if current_heading is not None:
                sections[current_heading].append(p.text)
            # else: paragraph appears before any heading — decide what to do

    return sections

def compare_documents(template_path, contract_path):
    """
    Compare a contract against a template.
    Returns a placeholder result for now — real comparison logic comes next.
    """
    template_text = extract_text(template_path)
    contract_text = extract_text(contract_path)

    print(f"Template has {len(template_text)} paragraphs.")
    print(f"Contract has {len(contract_text)} paragraphs.")

    # Template paragraphs
    # for i, p in enumerate(template_text):
    #     print(i, "->", p)

    # TODO: split into sections/clauses
    template_sections = split_into_sections(template_path)
    # contract_sections = split_into_sections(contract_path)

    print(f"Found {len(template_sections)} sections:")
    for heading, body in template_sections.items():
        print(f"  '{heading}' -> {len(body)} paragraphs")
    # Template sections
    # for i, p in enumerate(template_sections):
    #     print(i, "->", p)

    # TODO: align matching clauses between template and contract
    # TODO: flag missing / altered / unfilled clauses

    return {
        "status": "stub - comparison logic not yet implemented"
    }


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 3:
        print("Usage: python main.py <template.docx> <contract.docx>")
        sys.exit(1)

    result = compare_documents(sys.argv[1], sys.argv[2])
    print(result)